from flask import Flask, render_template, request,redirect,url_for,flash,send_from_directory,jsonify,flash
from wordcloud import WordCloud
from docx import Document
from pdf2docx import Converter
from docx2pdf import convert
from werkzeug.utils import secure_filename
from paddleocr import PaddleOCR,draw_ocr
import pandas as pd
import traceback,pdfkit
import erniebot
import pymysql
import win32api,win32con
import os,sys

## 初始化
app = Flask(__name__)
erniebot.api_type = 'aistudio'  
erniebot.access_token = 'c45a630b4d316eae8d412079a5c73685927aedba' 
app.secret_key = 'secret'  
app.config['SECRET_KEY'] = 'secret'  
UPLOAD_FOLDER = '../Smart_Editor_Web/uploads/'  
DOWNLOAD_FOLDER = '../Smart_Editor_Web/downloads'  
app.config['ALLOWED_EXTENSIONS'] = ['png','jpg', 'jpeg']

##检查是否存在文件夹
if not os.path.exists(UPLOAD_FOLDER):  
    os.makedirs(UPLOAD_FOLDER)  
if not os.path.exists(DOWNLOAD_FOLDER):  
    os.makedirs(DOWNLOAD_FOLDER) 

## 设置模型参数
sys.path.append(os.path.dirname(os.path.dirname(os.path.realpath(__file__))))   
folder_dir = os.path.dirname(os.path.realpath(__file__))
paddleOCR_dir = os.path.join(folder_dir, "PaddleOCR")
cls_model_dir="./PaddleOCR/inference/ch_ppocr_mobile_v2.0_cls_infer/" 
det_model_dir="./PaddleOCR/inference/ch_ppocr_mobile_v2.0_det_infer/"
rec_model_dir="./PaddleOCR/inference/ch_ppocr_mobile_v2.0_rec_infer/" 
use_angle_cls = True
use_space_char = True
use_gpu = False 
cls = True

## 登录界面
@app.route('/')
def login():
    return render_template('login.html')

@app.route('/login',methods=['GET','POST'])
def getLoginRequest():
    db = pymysql.connect(host="localhost", user="root", password="1023456zxc", database="user_data", charset="utf8")
    cursor = db.cursor()
    username = request.args.get('User_Name')
    password = request.args.get('User_Password')
    sql = "SELECT * FROM user_data.user_data WHERE User_Name=%s AND User_Password=%s"
    try:
        cursor.execute(sql, (username, password))
        results = cursor.fetchall()
        if len(results) == 1:
            db.close()
            return render_template('Functionality.html')  
        else:
            db.close()
            return '用户名或密码不正确',401
    except:
        traceback.print_exc()
        db.rollback()  
        db.close() 
        return jsonify({"error": "内部服务器错误"}), 500

## 注册界面
@app.route('/register')
def register():
    return render_template('register.html')

@app.route('/getRigistRequest')
def getRigistRequest():
    db = pymysql.connect(host="localhost", user="root", password="1023456zxc", database="user_data",charset="utf8") 
    cursor = db.cursor()
    username=request.args.get('User_Name')
    password1=request.args.get('User_Password')
    password2=request.args.get('User_Password')
    print(password1)
    print(password2)
    if password1==password2:
        sql = "INSERT INTO user_data.user(User_Name,User_Password) VALUES (%s ,%s)"
        try:
            cursor.execute(sql,(username,password1))
            db.commit()
            db.close()
            return render_template('login.html')
        except:
            traceback.print_exc()
            db.rollback()
            db.close() 
            win32api.MessageBox(0,"注册失败","提醒",win32con.MB_ICONWARNING)
            return render_template('register.html')
    else:
        win32api.MessageBox(0,"两次输入密码不一致，请重新输入！","提醒",win32con.MB_ICONWARNING)
        db.close()
        return render_template('register.html')
    
## 功能界面
@app.route('/Functionality',methods=['GET','POST'])
def Functionality():
    return render_template('Functionality.html')

## 文档格式转换界面
@app.route('/format_conversion',methods=['GET','POST'])  
def format_conversion():  
    return render_template('format_conversion.html')

# WORD转PDF
def allowed_file(filename):  
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ['docx','doc']

@app.route('/word_to_pdf', methods=['GET','POST'])
def word_to_pdf():
    return render_template('word_to_pdf.html')

@app.route('/wtp', methods=['GET','POST'])  
def wtp_page():  
    if 'file' not in request.files:  
        flash('No file part')  
        return '<script> alert("转换失败");</script>'   
    file = request.files['file']  
    if file.filename == '':  
        flash('No selected file')  
        return '<script> alert("转换失败");</script>'   
    if file and allowed_file(file.filename):  
        filename = file.filename  
        filepath = os.path.join(UPLOAD_FOLDER, filename)  
        file.save(filepath)  
      
        pdf_filename = os.path.splitext(filename)[0] + '.pdf'  
        pdf_filepath = os.path.join(DOWNLOAD_FOLDER, pdf_filename)  
        try:  
            convert(filepath, pdf_filepath)  
            return redirect('format_conversion') 
        except Exception as e:  
            flash(f'Error converting file: {e}')  
            return '<script> alert("转换失败");</script>' 
    else:  
        flash('Invalid file type. Only .docx files are allowed.')  
        return redirect('word_to_pdf')

# PDF转WORD
def convert_pdf_to_docx(pdf_path, docx_path):  
    cv = Converter(pdf_path)  
    cv.convert(docx_path, start=0, end=None)  
    cv.close() 

@app.route('/pdf_to_word', methods=['GET','POST'])
def pdf_to_word():
    return render_template('pdf_to_word.html')

@app.route('/ptw_page', methods=['GET','POST'])  
def ptw_page():  
    if 'file' not in request.files:  
        return 'No file part', 400  
    file = request.files['file']  
    if file.filename == '':  
        return 'No selected file', 400  
    if file:  
        temp_pdf_path = os.path.join(UPLOAD_FOLDER, file.filename)  
        file.save(temp_pdf_path)    
        temp_docx_path = os.path.join(DOWNLOAD_FOLDER, os.path.splitext(file.filename)[0] + '.docx')  
        if not os.path.exists(DOWNLOAD_FOLDER):
            os.makedirs(DOWNLOAD_FOLDER)
        try:  
            convert_pdf_to_docx(temp_pdf_path, temp_docx_path)
            message =f"转换成功"  
            return render_template( 'Functionality.html',message=message)
        except Exception as e:
            return f'Error converting {e}',500

# EXCEL转PDF
def allowed_file_excel(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ['xlsx']

@app.route('/excel_to_pdf', methods=['GET','POST'])
def excel_to_pdf():
    return render_template('excel_to_pdf.html')

@app.route('/etp_page',methods=['GET','POST'])
def etp_page():
    if request.method == 'POST':  
        if 'file' not in request.files:  
            flash('No file part')  
            return redirect(url_for('etp_page'))  
  
        file = request.files['file']  
        if file.filename == '':  
            flash('No selected file')  
            return redirect(url_for('etp_page'))  
  
        if file and allowed_file_excel(file.filename):  
            filename = secure_filename(file.filename)  
            filepath = os.path.join(UPLOAD_FOLDER, filename)  
            file.save(filepath)  
  
            pdf_filename = os.path.splitext(filename)[0] + '.pdf'  
            pdf_filepath = os.path.join(DOWNLOAD_FOLDER, pdf_filename)  
            
            try:  
                df = pd.read_excel(filepath)  
                html_str = df.to_html(index=False)  
                pdfkit.from_string(html_str, pdf_filepath)  
                flash('转换成功')
                return redirect('format_conversion') 
            except Exception as e:  
                print(f"Error converting Excel to PDF: {e}") 
                flash('Conversion failed')  
                return redirect(url_for('etp_page'))  
        else:  
            flash('File type not allowed')  
            return redirect(url_for('etp_page'))   
    return render_template('excel_to_pdf.html')


## 图片识别界面
def allowed_file_image(filename):
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/image_ocr_page', methods=['GET'])
def index():
    return render_template('image_ocr.html')

@app.route('/image_ocr', methods=['POST'])
def image_ocr(): 
    if 'file' not in request.files:  
        return jsonify({'error': 'No file part'}), 400  
    file = request.files['file']  
    if file.filename == '' or not allowed_file_image(file.filename):  
        return jsonify({'error': 'No selected file or invalid file type'}), 400  
    try:  
        filename = secure_filename(file.filename)  
        filepath = os.path.join(UPLOAD_FOLDER, filename)  
        file.save(filepath)  
        ocr = PaddleOCR(cls_model_dir=cls_model_dir, det_model_dir=det_model_dir, rec_model_dir=rec_model_dir, use_angle_cls=use_angle_cls, use_gpu=use_gpu )  
        result = ocr.ocr(filepath, cls=False) 
        doc = Document()  
        txts = []    
        for line in result:  
            for detection in line:  
                txts.append(detection[1][0])
        for i in range(len(txts)):
            doc.add_paragraph(txts[i])
        doc.save(os.path.join(DOWNLOAD_FOLDER, filename.capitalize() + '.docx'))   
        return '<script> alert("转换成功");</script>'  
    except Exception as e:  
        return jsonify({'error': str(e)}), 500 

## AI问答界面
@app.route('/ai_answer',methods=['GET','POST'])
def ai_answer():
    return render_template('ai_answer.html')

@app.route('/ai_answer_func',methods=['GET','POST'])
def ai_answer_func():    
    data = request.form.get('data', 'No data provided')  
    erniebot.api_type = 'aistudio'  
    erniebot.access_token = 'c45a630b4d316eae8d412079a5c73685927aedba'  
    try:   
        response = erniebot.ChatCompletion.create(  
            model='ernie-turbo',  
            messages=[{'role': 'user', 'content': data}]  
        )          
        result = response.get_result() 
        message = f"{result}"  
    except Exception as e:    
        message = f"An error occurred: {str(e)}"    
    return render_template('ai_answer.html',message=message) 

## 数据可视化界面
@app.route('/data_visualization')  
def data_visualization():  
    return render_template('data_visualization.html')
# 词云界面
def save_docx_as_txt(docx_path, txt_path):  
    doc = Document(docx_path)  
    with open(txt_path, 'w', encoding='utf-8') as txt_file:  
        for paragraph in doc.paragraphs:  
            txt_file.write(paragraph.text + '\n')

def allowed_file_docx(filename):  
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ['docx']

def generate_wordcloud(text):  
    # 可以自己制定词云样式 
    return WordCloud(width=800, height=400,font_path="msyh.ttc",background_color="white").generate(text)

def save_wordcloud(wordcloud, output_folder, filename):
    wordcloud_image_path = os.path.join(output_folder, filename)
    wordcloud.to_file(wordcloud_image_path)
    return wordcloud_image_path 

@app.route('/wordcloud_docx', methods=['POST', 'GET'])  
def wordcloud_docx():  
    if request.method == 'POST':  
        if 'file' not in request.files:  
            flash('No file part')  
            return redirect(request.url)  
        file = request.files['file']  
        if file.filename == '':  
            flash('No selected file')  
            return redirect(request.url)  
        if file and allowed_file_docx(file.filename):  
            filename = file.filename 
            docx_path = os.path.join(UPLOAD_FOLDER, filename)  
            file.save(docx_path)  
            txt_filename = os.path.splitext(filename)[0] + '.txt'  
            txt_path = os.path.join(UPLOAD_FOLDER, txt_filename)
  
            try:  
                save_docx_as_txt(docx_path, txt_path)  
                with open(txt_path, 'r', encoding='utf-8') as txt_file:  
                    text = txt_file.read() 
                wordcloud = generate_wordcloud(text)  
                wordcloud_filename = f"{os.path.splitext(filename)[0]}_wordcloud.png"  
                wordcloud_image_path = save_wordcloud(wordcloud, DOWNLOAD_FOLDER, wordcloud_filename)  
                flash(f'Wordcloud saved as {wordcloud_filename}')  
                return redirect(url_for('download_wordcloud', filename=wordcloud_filename))  

            except Exception as e:  
                flash(f'Error processing file: {e}')  
                return redirect(request.url)
    return render_template('Data_Visualization.html')

@app.route('/download_wordcloud/<filename>')  
def download_wordcloud(filename):  
    return send_from_directory(DOWNLOAD_FOLDER, filename)

## 信息抽取
# @app.route('/Information_Extraction',methods = ['GET','POST'])
# def Information_Extraction(filename):

#     return 

# 程序运行处
if __name__ == '__main__':
    app.run(debug=True)